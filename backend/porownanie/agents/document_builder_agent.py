import os
from docx import Document
from docx.shared import Pt
from ..models import ProtocolData

class DocumentBuilderAgent:
    def __init__(self, template_path: str = None):
        if template_path and os.path.exists(template_path) and os.path.getsize(template_path) > 0:
            self.document = Document(template_path)
        else:
            self.document = Document()
            self._init_default_structure()

    def _init_default_structure(self):
        self.document.add_heading('PROTOKÓŁ POSTĘPOWANIA', 0)
        p = self.document.add_paragraph('O UDZIELENIE ZAMÓWIENIA PUBLICZNEGO')
        p.alignment = 1 # Center

    def build_document(self, data: ProtocolData, output_path: str):
        # Add metadata
        self.document.add_paragraph(f"Numer sprawy: {data.case_id}")
        self.document.add_paragraph(f"Data: {data.creation_date.strftime('%Y-%m-%d')}")
        
        # Add sections
        sorted_sections = sorted(data.sections, key=lambda x: x.section_id)
        for section in sorted_sections:
            self.document.add_heading(f"{section.section_id}. {section.title}", level=1)
            self.document.add_paragraph(str(section.content))
            self.document.add_paragraph() # Spacer

        self.document.save(output_path)
        return output_path
